from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Penjelasan Kode Hasil Diagnosa CBR 79.41.docx"


def setup_document(document):
    section = document.sections[0]
    section.top_margin = Cm(4)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.bottom_margin = Cm(3)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True
    return paragraph


def add_paragraph(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_code(document, code):
    for line in code.strip().splitlines():
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)

    document.add_paragraph()
    return table


def main():
    document = Document()
    setup_document(document)

    add_heading(document, "Penjelasan Kode Hasil Diagnosa CBR", 1)
    add_paragraph(
        document,
        "Pada sistem pakar diagnosa kerusakan laptop gaming, hasil diagnosa "
        "ditentukan menggunakan metode Case-Based Reasoning (CBR). Metode ini "
        "bekerja dengan cara membandingkan kasus baru yang dimasukkan oleh user "
        "dengan basis kasus lama yang tersimpan di dalam sistem. Nilai kemiripan "
        "atau similarity dihitung berdasarkan bobot gejala yang cocok terhadap "
        "total bobot gejala pada setiap kasus kerusakan.",
    )

    add_heading(document, "1. Kode Perhitungan Similarity CBR", 2)
    add_paragraph(
        document,
        "Kode utama perhitungan similarity terdapat pada file src/lib/cbr.js. "
        "Fungsi yang digunakan adalah calculateSimilarity(). Fungsi ini menerima "
        "data basis kasus dan daftar gejala yang dipilih user, kemudian "
        "menghasilkan ranking kerusakan berdasarkan nilai similarity tertinggi.",
    )
    add_code(
        document,
        """
export function calculateSimilarity(db, selectedGejalaIds) {
  const selected = new Set(selectedGejalaIds.map(String));

  return db.kerusakan
    .map((kerusakan) => {
      const kasus = db.basisKasus.filter(
        (item) => item.idKerusakan === kerusakan.id
      );

      const totalBobot = kasus.reduce(
        (total, item) => total + Number(item.bobot),
        0
      );

      const bobotCocok = kasus.reduce((total, item) => {
        return selected.has(String(item.idGejala))
          ? total + Number(item.bobot)
          : total;
      }, 0);

      const similarity = totalBobot > 0 ? bobotCocok / totalBobot : 0;

      return {
        ...kerusakan,
        similarity,
        persentase: Number((similarity * 100).toFixed(2))
      };
    })
    .sort((a, b) => b.similarity - a.similarity);
}
""",
    )
    add_paragraph(
        document,
        "Pada kode tersebut, variabel selected digunakan untuk menyimpan daftar "
        "gejala yang dipilih oleh user. Sistem kemudian melakukan perulangan "
        "terhadap seluruh data kerusakan. Untuk setiap kerusakan, sistem mengambil "
        "basis kasus yang sesuai, menghitung total bobot, menghitung bobot gejala "
        "yang cocok, kemudian menghitung nilai similarity. Hasil akhirnya diurutkan "
        "berdasarkan nilai similarity tertinggi.",
    )

    add_heading(document, "2. Kode Penyimpanan Hasil Diagnosa", 2)
    add_paragraph(
        document,
        "Setelah nilai similarity dihitung, sistem mengambil data kerusakan dengan "
        "nilai similarity tertinggi sebagai hasil diagnosa utama. Data tersebut "
        "kemudian disimpan ke dalam riwayat diagnosa.",
    )
    add_code(
        document,
        """
const results = calculateSimilarity(db, selected);
const best = results[0];

const gejalaDipilih = db.gejala.filter((item) =>
  selected.includes(item.id)
);

const diagnosa = {
  id: nextId(db.diagnosa),
  namaUser,
  tanggal: new Date().toISOString(),
  status: 'pending',
  gejalaUtama,
  selectedGejala: selected,
  gejalaDipilih,
  catatan,
  gambar: imageResult?.path ?? null,
  idKerusakanHasil: best.id,
  hasilKerusakan: best.nama,
  nilaiSimilarity: best.similarity,
  results
};

db.diagnosa.unshift(diagnosa);
writeDb(db);
""",
    )
    add_paragraph(
        document,
        "Variabel results berisi daftar ranking kerusakan berdasarkan nilai "
        "similarity. Data pada indeks pertama, yaitu results[0], merupakan "
        "kerusakan dengan nilai similarity tertinggi. Nilai tersebut disimpan "
        "sebagai hasilKerusakan dan nilaiSimilarity pada data diagnosa.",
    )

    add_heading(document, "3. Kode Tampilan Hasil Diagnosa", 2)
    add_paragraph(
        document,
        "Pada halaman hasil diagnosa, sistem menampilkan nama kerusakan, tanggal "
        "diagnosa, nilai similarity dalam bentuk persentase, tingkat kerusakan, "
        "penjelasan, solusi, gejala yang dipilih, serta tabel perbandingan "
        "retrieve.",
    )
    add_code(
        document,
        """
<script>
  export let data;
  import { formatDate, formatPercent } from '$lib/format.js';

  const diagnosa = data.diagnosa;
  const best = diagnosa.results && diagnosa.results.length
    ? diagnosa.results[0]
    : null;
</script>

<div>
  <p>Hasil diagnosa untuk {diagnosa.namaUser}</p>
  <h1>{diagnosa.hasilKerusakan}</h1>
  <p>{formatDate(diagnosa.tanggal)}</p>
</div>

<div>
  <div>{formatPercent(diagnosa.nilaiSimilarity)}</div>
  {#if best}
    <span>Tingkat: {best.tingkat}</span>
  {/if}
</div>

<div
  class="h-full rounded-full bg-brand"
  style="width: {diagnosa.nilaiSimilarity * 100}%">
</div>
""",
    )
    add_paragraph(
        document,
        "Fungsi formatPercent() digunakan untuk mengubah nilai similarity yang "
        "berbentuk desimal menjadi persentase. Progress bar pada halaman hasil "
        "menggunakan nilai diagnosa.nilaiSimilarity dikalikan 100 sehingga panjang "
        "bar sesuai dengan tingkat kemiripan hasil diagnosa.",
    )

    add_heading(document, "4. Perhitungan Persentase 79.41%", 2)
    add_paragraph(
        document,
        "Pada hasil diagnosa yang ditampilkan, user memilih beberapa gejala, yaitu "
        "G01, G02, G03, G04, G06, dan G07. Berdasarkan proses retrieve, kerusakan "
        "dengan nilai similarity tertinggi adalah Overheating dengan nilai "
        "79.41%. Nilai tersebut diperoleh dari perbandingan bobot gejala yang cocok "
        "dengan total bobot kasus Overheating.",
    )

    add_table(
        document,
        ["Kode Gejala", "Nama Gejala", "Status"],
        [
            ["G01", "Laptop mati total", "Dipilih user"],
            ["G02", "Laptop cepat panas", "Dipilih user"],
            ["G03", "Kipas berbunyi keras", "Dipilih user"],
            ["G04", "FPS game menurun", "Dipilih user"],
            ["G06", "Layar blank hitam", "Dipilih user"],
            ["G07", "Keyboard tidak berfungsi", "Dipilih user"],
        ],
    )
    add_paragraph(
        document,
        "Tabel di atas menunjukkan gejala yang dipilih oleh user pada proses "
        "diagnosa. Gejala-gejala tersebut kemudian dibandingkan dengan basis kasus "
        "lama yang dimiliki oleh sistem.",
    )

    add_table(
        document,
        ["Kode Gejala", "Bobot", "Status Kecocokan"],
        [
            ["G02", "1.0", "Cocok"],
            ["G03", "0.9", "Cocok"],
            ["G04", "0.8", "Cocok"],
            ["G05", "0.7", "Tidak cocok"],
        ],
    )
    add_paragraph(
        document,
        "Tabel di atas menunjukkan basis kasus Overheating pada saat diagnosa "
        "dihitung. Dari basis kasus tersebut, gejala G02, G03, dan G04 cocok dengan "
        "gejala yang dipilih oleh user, sedangkan G05 tidak dipilih oleh user "
        "sehingga tidak dihitung sebagai bobot cocok.",
    )

    add_paragraph(document, "Perhitungan bobot gejala cocok adalah sebagai berikut:")
    add_code(
        document,
        """
Bobot gejala cocok = G02 + G03 + G04
                   = 1.0 + 0.9 + 0.8
                   = 2.7
""",
    )
    add_paragraph(document, "Perhitungan total bobot kasus Overheating adalah sebagai berikut:")
    add_code(
        document,
        """
Total bobot kasus Overheating = G02 + G03 + G04 + G05
                              = 1.0 + 0.9 + 0.8 + 0.7
                              = 3.4
""",
    )
    add_paragraph(document, "Dengan demikian, nilai similarity dihitung sebagai berikut:")
    add_code(
        document,
        """
Similarity = Bobot Cocok / Total Bobot
           = 2.7 / 3.4
           = 0.7941176470588235

Persentase = 0.7941176470588235 x 100
           = 79.41176470588235%
           = 79.41%
""",
    )
    add_paragraph(
        document,
        "Berdasarkan perhitungan tersebut, nilai similarity untuk kerusakan "
        "Overheating adalah 0.7941176470588235 atau 79.41%. Nilai ini menjadi hasil "
        "diagnosa utama karena memiliki similarity tertinggi dibandingkan "
        "kerusakan lainnya.",
    )

    add_heading(document, "5. Penjelasan Hasil Retrieve", 2)
    add_table(
        document,
        ["Kerusakan", "Similarity", "Tingkat"],
        [
            ["Overheating", "79.41%", "Sedang"],
            ["Kerusakan VGA", "78.13%", "Berat"],
            ["Kerusakan Power Supply", "40.00%", "Berat"],
            ["Kerusakan RAM", "28.00%", "Sedang"],
            ["Kerusakan Baterai", "0.00%", "Ringan"],
        ],
    )
    add_paragraph(
        document,
        "Tabel di atas menunjukkan hasil perbandingan retrieve dari seluruh data "
        "kerusakan. Sistem mengurutkan kerusakan berdasarkan nilai similarity "
        "tertinggi. Karena Overheating memiliki nilai similarity paling tinggi, "
        "yaitu 79.41%, maka sistem menetapkan Overheating sebagai hasil diagnosa "
        "utama.",
    )

    add_heading(document, "6. Catatan Mengenai Data Retain", 2)
    add_paragraph(
        document,
        "Pada data sistem saat ini, beberapa gejala tambahan seperti G01, G06, dan "
        "G07 dapat saja sudah masuk ke basis kasus Overheating karena adanya proses "
        "Retain setelah diagnosa divalidasi. Namun nilai 79.41% pada gambar "
        "merupakan hasil diagnosa yang sudah tersimpan sebelumnya, yaitu nilai yang "
        "dihitung pada saat proses diagnosa pertama kali dilakukan. Oleh karena itu, "
        "hasil historis diagnosa tidak berubah secara otomatis walaupun basis kasus "
        "bertambah setelah proses Retain.",
    )

    add_heading(document, "Kesimpulan", 2)
    add_paragraph(
        document,
        "Persentase 79.41% pada hasil diagnosa Overheating diperoleh dari "
        "perhitungan similarity CBR, yaitu jumlah bobot gejala yang cocok dibagi "
        "dengan total bobot kasus Overheating. Gejala yang cocok memiliki total "
        "bobot 2.7, sedangkan total bobot kasus Overheating adalah 3.4. Hasil "
        "perhitungan 2.7 / 3.4 menghasilkan nilai 0.7941176470588235 atau 79.41%. "
        "Nilai tersebut kemudian ditampilkan pada halaman hasil diagnosa sebagai "
        "persentase kemiripan kasus.",
    )

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
