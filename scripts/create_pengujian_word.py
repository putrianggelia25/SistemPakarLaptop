from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Pengujian Vitest dan Basic Path Testing CBR.docx"


def set_normal_style(document):
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.space_after = Pt(0)


def add_heading(document, text, level):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True
    return paragraph


def add_paragraph(document, text):
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return paragraph


def add_code(document, code):
    for line in code.strip().splitlines():
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
    document.add_paragraph()
    return table


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(4)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)
    section.bottom_margin = Cm(3)
    set_normal_style(document)

    add_heading(document, "5.5 Pengujian", 1)
    add_paragraph(
        document,
        "Pengujian sistem dilakukan untuk memverifikasi kelayakan fungsionalitas "
        "sistem pakar diagnosa kerusakan laptop gaming, khususnya pada penerapan "
        "metode Case-Based Reasoning (CBR). Pengujian difokuskan pada fungsi utama "
        "CBR, yaitu proses perhitungan nilai kemiripan kasus, validasi hasil "
        "diagnosa, revisi hasil diagnosa oleh admin, serta penyimpanan kasus baru "
        "ke dalam basis kasus.",
    )
    add_paragraph(
        document,
        "Pada penelitian ini, pengujian dilakukan menggunakan dua pendekatan utama, "
        "yaitu Pengujian Unit Menggunakan Vitest (Black-Box dan Functional Testing) "
        "serta Pengujian White-Box dengan Basic Path Testing. Pengujian unit "
        "digunakan untuk memverifikasi fungsi berdasarkan input dan output yang "
        "diharapkan, sedangkan basic path testing digunakan untuk menganalisis "
        "jalur logika internal dari metode CBR.",
    )

    add_heading(document, "5.5.1 Pengujian Unit Menggunakan Vitest", 2)
    add_paragraph(
        document,
        "Vitest merupakan framework pengujian unit modern pada JavaScript yang "
        "berjalan di atas Vite. Framework ini digunakan karena sistem dikembangkan "
        "menggunakan SvelteKit, Node.js, dan JavaScript, sehingga pengujian dapat "
        "dilakukan secara langsung terhadap fungsi utama sistem.",
    )
    add_paragraph(
        document,
        "Pengujian unit dilakukan pada modul CBR yang terdapat pada berkas "
        "src/lib/cbr.js. Fungsi yang diuji meliputi calculateSimilarity(), "
        "revise(), dan retain(). Konfigurasi perintah pengujian pada package.json "
        "adalah sebagai berikut:",
    )
    add_code(
        document,
        """
"scripts": {
  "test": "vitest run",
  "test:watch": "vitest"
}
""",
    )

    add_heading(document, "Kode Fungsi CBR yang Diuji", 3)
    add_paragraph(
        document,
        "Berikut adalah potongan kode utama metode CBR yang digunakan dalam sistem. "
        "Fungsi calculateSimilarity() digunakan untuk proses Retrieve dan Reuse, "
        "fungsi revise() digunakan untuk proses Revise, sedangkan fungsi retain() "
        "digunakan untuk proses Retain.",
    )
    cbr_code = (ROOT / "src/lib/cbr.js").read_text(encoding="utf-8")
    add_code(document, cbr_code)

    add_heading(document, "5.5.1.1 Pengujian Black-Box dan Functional Testing", 3)
    add_paragraph(
        document,
        "Pengujian black-box dilakukan untuk memastikan bahwa fungsi menghasilkan "
        "output sesuai dengan input yang diberikan. Pada pengujian ini, struktur "
        "internal program tidak diperhatikan. Fokus pengujian adalah kesesuaian "
        "hasil keluaran sistem terhadap kebutuhan fungsional.",
    )
    add_paragraph(
        document,
        "File pengujian black-box berada pada src/lib/cbr.blackbox.test.js. Contoh "
        "kode pengujian adalah sebagai berikut:",
    )
    blackbox_code = (ROOT / "src/lib/cbr.blackbox.test.js").read_text(encoding="utf-8")
    add_code(document, blackbox_code)

    add_heading(document, "Tabel 5.X Pengujian Black-Box dan Functional Testing", 4)
    add_table(
        document,
        ["No", "Kode Uji", "Fungsi yang Diuji", "Skenario Uji", "Input", "Ekspektasi Output", "Status"],
        [
            [1, "BB-01", "calculateSimilarity()", "Menghitung similarity dan ranking kerusakan", "Gejala [1, 2, 3]", "Ranking pertama Overheating, similarity 100%", "Lulus"],
            [2, "BB-02", "calculateSimilarity()", "Menghitung similarity parsial", "Gejala [1, 2]", "Similarity Overheating bernilai 72%", "Lulus"],
            [3, "BB-03", "calculateSimilarity()", "Gejala tidak cocok dengan basis kasus", "Gejala [99]", "Semua similarity bernilai 0", "Lulus"],
            [4, "BB-04", "revise()", "Admin menyetujui diagnosa", "Diagnosa ID 3, aksi setujui", "Status menjadi disetujui", "Lulus"],
            [5, "BB-05", "revise()", "Admin merevisi hasil diagnosa", "Revisi ke Kerusakan RAM", "Status menjadi direvisi dan hasil revisi tersimpan", "Lulus"],
            [6, "BB-06", "retain()", "Menolak retain sebelum validasi", "Diagnosa pending", "Muncul error validasi diagnosa terlebih dahulu", "Lulus"],
            [7, "BB-07", "retain()", "Menyimpan kasus valid ke basis kasus", "Diagnosa direvisi", "Kasus baru tersimpan dengan bobot 1 dan 0.8", "Lulus"],
        ],
    )
    add_paragraph(
        document,
        "Tabel di atas menunjukkan hasil pengujian black-box dan functional testing "
        "terhadap fungsi utama CBR. Seluruh skenario berhasil dijalankan sesuai "
        "dengan output yang diharapkan. Dengan demikian, fungsi CBR telah memenuhi "
        "kebutuhan fungsional sistem, terutama dalam menghitung similarity, "
        "melakukan revisi hasil, serta menyimpan kasus baru.",
    )

    add_heading(document, "5.5.2 Pengujian White-Box Dengan Basic Path Testing", 2)
    add_paragraph(
        document,
        "Pengujian white-box dilakukan dengan memperhatikan struktur logika internal "
        "program. Metode Basic Path Testing digunakan untuk memastikan bahwa setiap "
        "jalur independen pada fungsi CBR telah diuji minimal satu kali.",
    )
    add_paragraph(
        document,
        "Pengujian ini dilakukan terhadap tiga fungsi utama, yaitu "
        "calculateSimilarity(), revise(), dan retain(). File pengujian white-box "
        "berada pada src/lib/cbr.whitebox.test.js.",
    )
    whitebox_code = (ROOT / "src/lib/cbr.whitebox.test.js").read_text(encoding="utf-8")
    add_code(document, whitebox_code)

    add_heading(document, "5.5.2.1 Analisis White-Box Fungsi calculateSimilarity()", 3)
    add_paragraph(
        document,
        "Fungsi calculateSimilarity() memiliki beberapa kondisi logika utama, yaitu "
        "membentuk kumpulan gejala yang dipilih user, melakukan perulangan data "
        "kerusakan, mengambil basis kasus berdasarkan kerusakan, menghitung total "
        "bobot, menghitung bobot gejala yang cocok, mengecek kondisi totalBobot > 0, "
        "menghasilkan nilai similarity, dan mengurutkan hasil berdasarkan nilai "
        "similarity tertinggi.",
    )
    add_table(
        document,
        ["Node", "Keterangan"],
        [
            [1, "Titik masuk fungsi calculateSimilarity()"],
            [2, "Membentuk kumpulan gejala yang dipilih user"],
            [3, "Melakukan perulangan data kerusakan"],
            [4, "Mengambil basis kasus berdasarkan kerusakan"],
            [5, "Menghitung total bobot"],
            [6, "Menghitung bobot gejala yang cocok"],
            [7, "Mengecek kondisi totalBobot > 0"],
            [8, "Menghasilkan nilai similarity"],
            [9, "Mengurutkan hasil berdasarkan similarity"],
            [10, "Titik keluar fungsi"],
        ],
    )
    add_table(
        document,
        ["Jalur", "Alur Node", "Keterangan"],
        [
            ["Jalur 1", "1-2-3-4-5-6-7-8-9-10", "Terdapat basis kasus dan gejala cocok"],
            ["Jalur 2", "1-2-3-4-5-6-7-8-9-10", "Terdapat basis kasus tetapi gejala tidak cocok"],
            ["Jalur 3", "1-2-3-4-5-7-8-9-10", "Tidak terdapat basis kasus sehingga similarity bernilai 0"],
        ],
    )

    add_heading(document, "5.5.2.2 Analisis White-Box Fungsi revise()", 3)
    add_paragraph(
        document,
        "Fungsi revise() memiliki percabangan utama berdasarkan aksi admin, yaitu "
        "setujui, revisi, dan aksi tidak dikenali. Fungsi ini juga memiliki jalur "
        "validasi ketika data diagnosa tidak ditemukan, ID kerusakan revisi tidak "
        "valid, dan kerusakan revisi tidak ditemukan.",
    )
    add_table(
        document,
        ["Node", "Keterangan"],
        [
            [1, "Titik masuk fungsi revise()"],
            [2, "Mencari data diagnosa berdasarkan ID"],
            [3, "Mengecek apakah diagnosa ditemukan"],
            [4, "Mengecek aksi setujui"],
            [5, "Mengubah status menjadi disetujui"],
            [6, "Mengecek aksi revisi"],
            [7, "Mengecek ID kerusakan revisi"],
            [8, "Mencari data kerusakan revisi"],
            [9, "Mengubah status menjadi direvisi"],
            [10, "Menampilkan error jika aksi tidak dikenali"],
            [11, "Titik keluar fungsi"],
        ],
    )
    add_table(
        document,
        ["Jalur", "Alur Node", "Keterangan"],
        [
            ["Jalur 1", "1-2-3-4-5-11", "Diagnosa ditemukan dan aksi setujui"],
            ["Jalur 2", "1-2-3-4-6-7-8-9-11", "Diagnosa ditemukan dan aksi revisi valid"],
            ["Jalur 3", "1-2-3-11", "Diagnosa tidak ditemukan"],
            ["Jalur 4", "1-2-3-4-6-10-11", "Aksi tidak dikenali"],
            ["Jalur 5", "1-2-3-4-6-7-11", "ID kerusakan revisi tidak valid"],
            ["Jalur 6", "1-2-3-4-6-7-8-11", "Kerusakan revisi tidak ditemukan"],
        ],
    )

    add_heading(document, "5.5.2.3 Analisis White-Box Fungsi retain()", 3)
    add_paragraph(
        document,
        "Fungsi retain() memiliki struktur logika yang lebih kompleks karena "
        "menentukan apakah suatu kasus dapat disimpan ke basis kasus atau tidak. "
        "Fungsi ini memeriksa data diagnosa, status validasi, status retained, "
        "kerusakan final, serta pasangan kerusakan dan gejala yang akan disimpan.",
    )
    add_table(
        document,
        ["Node", "Keterangan"],
        [
            [1, "Titik masuk fungsi retain()"],
            [2, "Mencari data diagnosa"],
            [3, "Mengecek apakah diagnosa ditemukan"],
            [4, "Mengecek status validasi diagnosa"],
            [5, "Mengecek apakah kasus sudah pernah di-retain"],
            [6, "Menentukan kerusakan final dari hasil revisi"],
            [7, "Menentukan kerusakan final dari ranking pertama"],
            [8, "Fallback mencari kerusakan berdasarkan nama"],
            [9, "Mengecek apakah kerusakan final ditemukan"],
            [10, "Melakukan perulangan gejala terpilih"],
            [11, "Mengecek apakah pasangan kerusakan-gejala sudah ada"],
            [12, "Menambahkan kasus baru ke basis kasus"],
            [13, "Menandai diagnosa sebagai retained"],
            [14, "Titik keluar fungsi"],
        ],
    )
    add_table(
        document,
        ["Jalur", "Alur Node", "Keterangan"],
        [
            ["Jalur 1", "1-2-3-14", "Diagnosa tidak ditemukan"],
            ["Jalur 2", "1-2-3-4-14", "Diagnosa belum divalidasi"],
            ["Jalur 3", "1-2-3-4-5-14", "Kasus sudah pernah di-retain"],
            ["Jalur 4", "1-2-3-4-5-6-10-11-12-13-14", "Retain menggunakan hasil revisi"],
            ["Jalur 5", "1-2-3-4-5-7-10-11-13-14", "Retain menggunakan hasil retrieve tertinggi"],
            ["Jalur 6", "1-2-3-4-5-8-10-11-12-13-14", "Retain menggunakan fallback nama kerusakan"],
            ["Jalur 7", "1-2-3-4-5-8-9-14", "Kerusakan final tidak ditemukan"],
            ["Jalur 8", "1-2-3-4-5-6-10-11-13-14", "Data kasus sudah ada sehingga tidak ditambahkan ulang"],
        ],
    )

    add_heading(document, "Tabel 5.X Pengujian White-Box Basic Path", 4)
    add_table(
        document,
        ["No", "Kode Uji", "Fungsi", "Jalur yang Diuji", "Kondisi Pengujian", "Ekspektasi Output", "Status"],
        [
            [1, "WB-01", "calculateSimilarity()", "Jalur total bobot > 0", "Gejala [1] cocok", "Similarity dihitung sesuai bobot", "Lulus"],
            [2, "WB-02", "calculateSimilarity()", "Jalur total bobot = 0", "Kerusakan tanpa basis kasus", "Similarity bernilai 0", "Lulus"],
            [3, "WB-03", "revise()", "Diagnosa tidak ditemukan", "ID diagnosa tidak tersedia", "Error Diagnosa tidak ditemukan", "Lulus"],
            [4, "WB-04", "revise()", "Aksi tidak dikenali", "Aksi hapus", "Error Aksi revisi tidak dikenali", "Lulus"],
            [5, "WB-05", "revise()", "ID revisi tidak valid", "ID kerusakan 0", "Error Pilih kerusakan revisi yang valid", "Lulus"],
            [6, "WB-06", "revise()", "Kerusakan revisi tidak ditemukan", "ID kerusakan 999", "Error Kerusakan revisi tidak ditemukan", "Lulus"],
            [7, "WB-07", "retain()", "Diagnosa tidak ditemukan", "ID diagnosa tidak tersedia", "Error Diagnosa tidak ditemukan", "Lulus"],
            [8, "WB-08", "retain()", "Kasus sudah di-retain", "Status retained = true", "Tidak menambah data, added = 0", "Lulus"],
            [9, "WB-09", "retain()", "Menggunakan hasil retrieve", "Diagnosa disetujui tanpa revisi", "Ranking pertama digunakan sebagai kerusakan final", "Lulus"],
            [10, "WB-10", "retain()", "Fallback nama kerusakan", "results kosong", "Kerusakan dicari berdasarkan nama", "Lulus"],
            [11, "WB-11", "retain()", "Kerusakan final tidak ditemukan", "Nama kerusakan tidak tersedia", "Error Tidak ada kerusakan yang dapat di-retain", "Lulus"],
        ],
    )
    add_paragraph(
        document,
        "Tabel di atas menunjukkan bahwa setiap jalur logika penting pada fungsi "
        "CBR telah diuji. Pengujian mencakup jalur normal, jalur validasi, jalur "
        "revisi, jalur penyimpanan kasus, serta jalur kesalahan. Dengan demikian, "
        "struktur internal metode CBR dapat dikatakan berjalan sesuai rancangan.",
    )

    add_heading(document, "5.5.3 Hasil Pengujian", 2)
    add_paragraph(
        document,
        "Pengujian dijalankan menggunakan perintah npm run test. Hasil eksekusi "
        "pengujian menunjukkan bahwa dua file pengujian berhasil dijalankan dengan "
        "total 18 skenario pengujian lulus.",
    )
    add_code(
        document,
        """
Test Files  2 passed (2)
Tests       18 passed (18)
""",
    )
    add_paragraph(
        document,
        "Selain itu, dilakukan validasi tambahan menggunakan npm run check dan "
        "npm run build. Hasil validasi menunjukkan bahwa project tidak memiliki "
        "error dan dapat dibangun dengan baik.",
    )
    add_code(
        document,
        """
npm run check  : 0 error, 0 warning
npm run build  : berhasil
""",
    )
    add_heading(document, "Tabel 5.X Rangkuman Hasil Pengujian", 4)
    add_table(
        document,
        ["Jenis Pengujian", "Jumlah Skenario", "Berhasil", "Gagal"],
        [
            ["Black-Box dan Functional Testing", 7, 7, 0],
            ["White-Box Basic Path Testing", 11, 11, 0],
            ["Total", 18, 18, 0],
        ],
    )
    add_paragraph(
        document,
        "Berdasarkan tabel rangkuman hasil pengujian, seluruh skenario pengujian "
        "berhasil dijalankan. Dari total 18 skenario, sebanyak 7 skenario merupakan "
        "pengujian black-box dan functional testing, sedangkan 11 skenario merupakan "
        "pengujian white-box basic path testing.",
    )
    add_paragraph(
        document,
        "Dengan hasil tersebut, dapat disimpulkan bahwa metode Case-Based Reasoning "
        "(CBR) pada sistem pakar diagnosa kerusakan laptop gaming telah berjalan "
        "sesuai rancangan. Fungsi perhitungan similarity, validasi hasil diagnosa, "
        "revisi hasil, dan penyimpanan kasus baru ke basis kasus telah menghasilkan "
        "output yang sesuai dengan kebutuhan sistem.",
    )

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
